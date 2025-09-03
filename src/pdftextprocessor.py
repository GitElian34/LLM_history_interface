import os
from sqlite3 import IntegrityError

import requests
import re
import tkinter as tk

from docx import Document
from typing import Optional
from tkinter import filedialog
from src.database.db_insert import *
from src.pdftest import process_pdf
from flairtest import findREN
from src.extract_text import *

class PDFTextProcessor:
    def __init__(self):
        """Initialise le processeur de texte PDF"""
        self.llms = ["llama3","gemma","llama2","mistral"] #Toutes les LLMs utilisé via ollama
        self.currentdates = {}  #variable pour stocker les années que donnent le LLM
        self.periode_histo = []  #variable pour stocker les périodes que donnent le LLM sur le texte
        self.periode_histo_ftn = []   #variable pour stocker les périodes que donnent le LLM sur les footnotes
        self.ftn_year = {}      #variable pour stocker les années que donnent le LLM pour les footnotes
        self.data_clear=[]  #currentdates mais passé par la fonctiion de seuil (qu'on utilise pas vraiment au final)

        self.entities_by_type = { #Hashmap pour stocker les Entités nomées de typer person, org et localisation
        'PERSON': [],
        'LOC': [],
        'ORG': []
    }



        self.content = None #contenu des articles privés des footnotes
        self.foot_notes = None #footnotes de l'article

    def ask_ollama(self, prompt: str, model: str ) -> Optional[str]:
        """
        Envoie une requête à l'API Ollama.

        Args:
            prompt: Prompt à envoyer
            model: Modèle à utiliser

        Returns:
            str: Réponse de l'API si succès, None sinon
        """
        try:
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={"model": model, "prompt": prompt, "stream": False}
            )
            response.raise_for_status()
            return response.json().get("response")
        except requests.exceptions.RequestException as e:
            print(f"Erreur lors de la requête à Ollama: {str(e)}")
            return None
        except Exception as e:
            print(f"Erreur inattendue: {str(e)}")
            return None


    def ask_question(self,question: str,content: str, llm :str) -> Optional[str]:

        full_prompt = f"{content}\n\n{question}"
        return self.ask_ollama(full_prompt, llm)

    import re
    from typing import List, Union, Tuple

    def extraire_annees_historiques(self, texte: str,epoques) -> List[Union[int, Tuple[int, int]]]:
        """
        Extrait les années historiques (3 ou 4 chiffres, plages d'années, et siècles '12th', '13th century', etc.)
        Retourne une liste sans doublons avec des entiers pour les années simples
        et des tuples pour les plages (début, fin).
        """
        result = []
        seen = set()

        roman_to_int = {
            'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
            'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10,
            'XI': 11, 'XII': 12, 'XIII': 13, 'XIV': 14, 'XV': 15,
            'XVI': 16, 'XVII': 17, 'XVIII': 18, 'XIX': 19, 'XX': 20
        }
        matches = re.findall(r'\b(\d{3,4})\b', texte)
        h_matches=re.findall( r'\[([A-ZÀ-ÿ][^\]]+)\]',texte)
        for nb in matches:
            try:
                num = int(nb)
                if num not in seen:  # Évite les doublons
                    result.append(num)
                    seen.add(num)
            except ValueError:  # Sécurité si conversion échoue
                print(f"Valeur numérique invalide : {nb}")

        # Traitement des périodes historiques
        for h_ep in h_matches:
            if h_ep not in epoques:  # Évite les doublons
                print(f"Période historique trouvée : {h_ep}")  # Debug optionnel
                epoques.append(h_ep)



        # 1. Extraction des siècles numériques (12th, 13th century/siècle)
        numeric_pattern = r'\b(\d{1,2})(?:st|nd|rd|th|e)\s*(?:century|siècle)?\b'
        for century_str in re.findall(numeric_pattern, texte, re.IGNORECASE):
            century = int(century_str)
            debut = (century - 1) * 100
            fin = century * 100
            if (debut, fin) not in seen:
                result.append((debut, fin))
                seen.add((debut, fin))

        # 2. Extraction des chiffres romains (XV, XII siècle/century)
        roman_pattern = r'\b(X{0,3}I{0,3}X?I{0,3})\s*(?:century|siècle|e)\b'
        for roman in re.findall(roman_pattern, texte, re.IGNORECASE):
            if roman.upper() in roman_to_int:
                century = roman_to_int[roman.upper()]
                debut = (century - 1) * 100
                fin = century * 100
                if (debut, fin) not in seen:
                    result.append((debut, fin))
                    seen.add((debut, fin))



        return sorted(result, key=lambda x: x[0] if isinstance(x, tuple) else x)

    def compter_occurrences_manuel(self, nombres, list):
        """Permet de compléter la hashmap de compter combien de fois chaque nombre est évoqué si on lance plusieurs requêtes au LLM
           sur un même article"""
        for nombre in nombres:
            if nombre in list:
                list[nombre] += 1
            else:
                list[nombre] = 1
        return list

    def clear_data(self, seuil: int, dico) -> List[Union[int, Tuple[int, int]]]:
        """Continuité de la fonction précédente, cette fonction permet de placer un seuil si on a lancé plusieurs prompts
         ou plusieurs fois le même prompt afin d'affiner les résultats et empêcher les hallucinations"""
        return [key for key, count in dico.items() if count >= seuil]

    def stringtoword(self,text: str) :
        """  Crée un document Word (.docx) à partir d'une chaîne de texte.
    Le texte est ajouté dans un paragraphe, puis sauvegardé dans un fichier.
    Cela permet d'extraire de façon optimal les footnotes"""
        doc = Document()
        p = doc.add_paragraph()
        p.add_run(text)

        # Ajouter du texte en grande taille

        doc.save("C:/Users/elian/Documents/stage/Recherche/output/output.docx")
        return doc

    def dict_to_text(self,dictionary):
        text = ""
        for page, content in dictionary.items():
            text += f"{page}:\n{content}\n\n"  # \n\n pour saut de ligne double entre pages
        return text.strip()  # .strip() pour supprimer le dernier \n inutile



    def extract_and_store(self, dict_, nombres_cibles, entities_by_type, epoques_cible: List[str], article_id, pattern=None):
        """
        NON UTILISE POUR L'INSTANT


        Parcourt les textes dans le dictionnaire et insère en BDD les entités trouvées
        (PERSON, LOC, ORG, nombres, époques), sans modifier le texte.

        Args:
            dict_: Dictionnaire contenant les valeurs à analyser
            nombres_cibles: Liste de nombres à détecter
            entities_by_type: Dictionnaire contenant les entités classées par type (PERSON, LOC, ORG)
            epoques_cible: Liste d'époques à détecter
            article_id: ID de l'article auquel rattacher les entités
            pattern: Optionnel - motif regex prédéfini (non utilisé ici)

        Returns:
            None (insertion uniquement en BDD)
        """
        # Convertir les nombres en strings une fois pour toutes
        nombres_str = list(map(str, nombres_cibles))

        # Traiter chaque valeur du dictionnaire
        for key, value in dict_.items():
            current_text = str(value)

            # === PERSON ===
            for entity in entities_by_type.get("PERSON", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                if re.search(pattern, current_text, flags=re.IGNORECASE):
                    insert_item(article_id, cleaned, "PERSON", "findREN")

            # === LOC ===
            for entity in entities_by_type.get("LOC", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                if re.search(pattern, current_text, flags=re.IGNORECASE):
                    insert_item(article_id, cleaned, "LOC", "findREN")

            # === ORG ===
            for entity in entities_by_type.get("ORG", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                if re.search(pattern, current_text, flags=re.IGNORECASE):
                    insert_item(article_id, cleaned, "ORG", "findREN")

            # === Nombres ===
            for num in nombres_str:
                pattern = r'(?<!\w)(' + re.escape(num) + r')(?!\w)'
                if re.search(pattern, current_text):
                    insert_item(article_id, num, "Number", "LLM")

            # === Époques ===
            for epoque in epoques_cible:
                cleaned = epoque.replace(" ", "_").replace("-", "_")

                pattern = r'(?<!\w)(' + re.escape(epoque) + r')(?!\w)'
                if re.search(pattern, current_text, flags=re.IGNORECASE):
                    insert_item(article_id, cleaned, "Period", "LLM")

    def number_to_bold_noDB(self, dict, nombres_cibles, entities_by_type, epoques_cible: List[str], pattern=None):
        """
         NON UTILISE POUR L'INSTANT





        Identique à number_to_bold mais sans insertion en base de données.
        Met en évidence nombres, entités et époques dans les valeurs du dictionnaire.

        Args:
            dict: Dictionnaire contenant les valeurs à modifier
            nombres_cibles: Liste de nombres à mettre en évidence
            entities_by_type: Dictionnaire contenant les entités classées par type (PERSON, LOC, ORG)
            epoques_cible: Liste d'époques à mettre en évidence
            pattern: Optionnel - motif regex prédéfini (non utilisé dans cette version)

        Returns:
            Le dictionnaire modifié
        """
        # Convertir les nombres en strings une fois pour toutes
        nombres_str = list(map(str, nombres_cibles))

        # Traiter chaque valeur du dictionnaire
        for key, value in dict.items():
            current_text = str(value)  # Conversion en string pour sécurité

            # === PERSON ===
            for entity in entities_by_type.get("PERSON", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_person(match):
                    found = match.group(1)
                    return f'++{cleaned}++'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_person, current_text, flags=re.IGNORECASE)

            # === LOC ===
            for entity in entities_by_type.get("LOC", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_loc(match):
                    found = match.group(1)
                    return f'@@{cleaned}@@'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_loc, current_text, flags=re.IGNORECASE)

            # === ORG ===
            for entity in entities_by_type.get("ORG", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_org(match):
                    found = match.group(1)
                    return f'┤┤{cleaned}┤┤'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_org, current_text, flags=re.IGNORECASE)

            # === Nombres ===
            for num in nombres_str:
                pattern = r'(?<!\w)(' + re.escape(num) + r')(?!\w)'

                def replacer_num(match):
                    found = match.group(1)
                    return f'**{found}**'

                current_text = re.sub(pattern, replacer_num, current_text)

            # === Époques ===
            for epoque in epoques_cible:
                cleaned = epoque.replace(" ", "_").replace("-", "_")

                def replacer_epoque(match):
                    found = match.group(1)
                    return f'¦¦{cleaned}¦¦'

                pattern = r'(?<!\w)(' + re.escape(epoque) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_epoque, current_text, flags=re.IGNORECASE)

            # Mettre à jour la valeur dans le dict
            dict[key] = current_text

        return dict

    def number_to_bold_withDB(self, dict, article_id: int, nombres_cibles, entities_by_type, epoques_cible: List[str],
                              pattern=None):
        """
        Met en évidence nombres, entités et époques dans les valeurs du dictionnaire
        ET insère chaque élément dans la base de données.

        Args:
            dict: Dictionnaire contenant les valeurs à modifier
            article_id: ID de l'article pour insertion en BDD
            nombres_cibles: Liste de nombres à mettre en évidence
            entities_by_type: Dictionnaire contenant les entités classées par type (PERSON, LOC, ORG)
            epoques_cible: Liste d'époques à mettre en évidence
            pattern: Optionnel - motif regex prédéfini (non utilisé ici)

        Returns:
            Le dictionnaire modifié
        """
        # Convertir les nombres en strings une fois pour toutes
        nombres_str = list(map(str, nombres_cibles))

        for key, value in dict.items():
            current_text = str(value)  # Sécurité

            # === PERSON ===
            for entity in entities_by_type.get("PERSON", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_person(match):
                    insert_item(article_id, cleaned, "PERSON", "highlight")  # insertion BDD
                    return f'++{cleaned}++'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_person, current_text, flags=re.IGNORECASE)

            # === LOC ===
            for entity in entities_by_type.get("LOC", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_loc(match):
                    insert_item(article_id, cleaned, "LOC", "highlight")
                    return f'@@{cleaned}@@'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_loc, current_text, flags=re.IGNORECASE)

            # === ORG ===
            for entity in entities_by_type.get("ORG", []):
                cleaned = entity.replace(" ", "_").replace("-", "_").replace(".", "\u00B7").replace("'", "\u02BC")

                def replacer_org(match):
                    insert_item(article_id, cleaned, "ORG", "highlight")
                    return f'┤┤{cleaned}┤┤'

                pattern = r'(?<!\w)(' + re.escape(entity) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_org, current_text, flags=re.IGNORECASE)

            # === Nombres ===
            for num in nombres_str:
                pattern = r'(?<!\w)(' + re.escape(num) + r')(?!\w)'

                def replacer_num(match):
                    found = match.group(1)
                    insert_item(article_id, found, "Number", "highlight")
                    return f'**{found}**'

                current_text = re.sub(pattern, replacer_num, current_text)

            # === Époques ===
            for epoque in epoques_cible:
                cleaned = epoque.replace(" ", "_").replace("-", "_")

                def replacer_epoque(match):
                    found = match.group(1)
                    insert_item(article_id, found, "Period", "highlight")
                    return f'¦¦{cleaned}¦¦'

                pattern = r'(?<!\w)(' + re.escape(epoque) + r')(?!\w)'
                current_text = re.sub(pattern, replacer_epoque, current_text, flags=re.IGNORECASE)

            # Mise à jour de la valeur dans le dict
            dict[key] = current_text

        return dict

    def choisir_dossier(self):
        """Ouvre une fenêtre de sélection de dossier et retourne son chemin"""
        root = tk.Tk()
        root.update_idletasks()
        root.lift()
        root.attributes('-topmost', True)
        root.focus_force()

        folder_path = filedialog.askdirectory(title="Choisissez un dossier", parent=root)

        root.destroy()
        return folder_path
    def get_all_pdfs(self,folder_path: str):
        """
        Parcourt un dossier et retourne la liste des fichiers PDF (sans extension).
        """
        pdf_files = []
        for file in os.listdir(folder_path):
            if file.lower().endswith(".pdf"):
                # on enlève l'extension pour ne garder que le "nom"
                pdf_files.append(f"{folder_path}/{os.path.splitext(file)[0]}.pdf")
        return pdf_files

    def extract_article_ids(self, article_list):
        """Extrait uniquement les IDs (après le dernier '_') d'une liste d'articles, sans le .pdf"""
        ids = []
        for article in article_list:
            # Prend la partie après le dernier underscore
            article_id = article.split("_")[-1]
            # Supprime l'extension .pdf si elle est présente
            article_id = article_id.replace(".pdf", "")
            ids.append(article_id)
        return ids

    def extract_article_names(self, article_list):
        """Extrait le nom entier d'une liste d'articles, sans le .pdf"""
        names = []
        for article in article_list:
            # Supprime l'extension .pdf si elle est présente

            article_name=article.replace(".pdf", "")
            article_name = article_name.split("/")[-1]
            names.append(article_name)
        return names

    def getdataLLM(self,iter:int,contenu , nbllm:int,epoques, question: str, result ):
        """
           Analyse un document PDF avec plusieurs modèles de LLM.

           Arguments :
               contexte (str) : Objectif général de l'analyse (ex. mettre en avant les dates historiques). Optionnel ici
               mais suffit d'être rajouté dans le prompt

               prompt (str) : Question envoyée aux modèles.
               nbllm (int) : Nombre de modèles LLM différent que l'on veut utiliser (entre 1 et 4).
               iter (int) : Nombre d’itérations à effectuer par modèle.

           Fonctionnement :
               - Parcourt chaque page du contenu du PDF. Cela permet de préciser les réponses en ne donnant au LLM pas trop de contenu d'un coup
               - Envoie la page et le prompt au modèle choisi.
               - Extrait et compte les dates/périodes historiques trouvées.
               - Applique une reconnaissance d’entités nommées (REN).
               - Répète l’opération pour tous les modèles et itérations.
           """


        contexte = ("Mon objectif est de prendre un pdf et de mettre en avant toutes"
                    "les dates et périodes historiques en les mettants en gras et plus gros"
                    "d'où l'importace que tu me mettes uniquement ce que tu vois et pas ce que tu déduis \n")
        prompt = question
        for i in range(0, nbllm):
            for j in range(0, iter):
                for page in contenu.values():
                    response = self.ask_question(prompt, page, self.llms[i])
                    if response:
                        print("Réponse reçue:")
                        annees_trouve = self.extraire_annees_historiques(response,epoques)
                        self.compter_occurrences_manuel(annees_trouve, result)
                        if type(page) != list:
                            findREN(page, 0.80,self.entities_by_type)


                    else:
                        print("Le traitement a échoué.")
        print("Fin du traitement")


    def FillDB(self, iter:int, nbllm:int,seuil :int, question: str,input_path: str,article_id: int ):
        """
          Traite un article PDF et insère son contenu enrichi dans la base de données.

          Arguments :
              iter (int) : Nombre d’itérations de traitement à effectuer.
              nbllm (int) : Nombre de modèles LLM utilisés pour l’analyse.
              seuil (int) : Seuil appliqué pour filtrer/clarifier les données extraites.
              question (str) : Prompt envoyé aux modèles pour guider l’extraction.
              input_path (str) : Chemin du fichier PDF d’entrée.
              article_id (int) : Identifiant de l’article associé dans la base de données.

          Fonctionnement :
              - Convertit le PDF en texte et en notes de bas de page.
              - Extrait et nettoie les données avec les LLM.
              - Filtre les résultats selon le seuil.
              - Traite les notes de bas de page séparément.
              - Supprime les notes du texte principal.
              - Inssère toutes les entités nommées ( dates périodes localisation personnes et organisation) dans la BDD
              - Insère les pages enrichies (dates/entités mises en avant) dans la base.
          """


        print("DEBUT DU TEST")
        #content = self.pdftotext.clean_content
        self.content,self.foot_notes= process_pdf(input_path)

        #print(self.content)
        self.getdataLLM(iter,self.content,nbllm,self.periode_histo,question,self.currentdates)
        self.data_clear = self.clear_data(seuil,self.currentdates)

        self.getdataLLM(iter,self.foot_notes,1,self.periode_histo_ftn,question,self.ftn_year)


        self.texte_clean = text_wo_footnotes(self.foot_notes,extraire_texte_pdf_par_page(input_path))
        insert_pages_dict(article_id,self.number_to_bold_withDB(self.texte_clean, article_id,self.data_clear, self.entities_by_type, self.periode_histo))






if __name__ == "__main__":
    print("début du programme")
    processor = PDFTextProcessor()
    question= """Relève simplement toutes les années et période historique de cet article scientifique, cela peut être
    des années comme des Noms propres ex: Moyen Age,grande depression...  N'invente RIEN, uniquement des dates et périodes que je
    pourraient voir écrite dans l'article. Mets les période historique entre crochet exemple :[Moyen Age]"""
#Permet à l'utilissateur de remplir la BDD avec les articles qu'il veut dans le dossier qu'il choisit
    choisir_directory= processor.choisir_dossier()
    all_article_path=processor.get_all_pdfs(choisir_directory)
    all_article_ids= processor.extract_article_ids(all_article_path)  #Celui là sert pour mes propres dossiers des articles de Persee afin que ça soit plus propre
    all_article_names=processor.extract_article_names(all_article_path)

    for i in range(0,len(all_article_ids)): #parcours tous les pdfs pour les mettre dans la BDD
        print("On passe à l'article :")
        print(all_article_names[i])
        processor.FillDB(1, 1, 0, question, all_article_path[i], all_article_names[i])
